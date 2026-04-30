"""Tiny verification tests for each fixture builder.

Each test:
- builds to tmp_path
- asserts file exists and has non-zero size
- checks basic structural integrity (e.g. opens with the right library)
"""

import zipfile

import openpyxl
from bs4 import BeautifulSoup
from docx import Document

from tests.fixtures.builders import (
    build_auto_cam_xlsx,
    build_bank_statement_pdf,
    build_case_zip,
    build_checklist_xlsx,
    build_dedupe_xlsx,
    build_equifax_html,
    build_pd_sheet_docx,
)


class TestAutoCamBuilder:
    def test_creates_file_with_four_sheets(self, tmp_path):
        out = build_auto_cam_xlsx(tmp_path / "auto_cam.xlsx")
        assert out.exists()
        assert out.stat().st_size > 0
        wb = openpyxl.load_workbook(out)
        assert wb.sheetnames == ["SystemCam", "Elegibilty", "CM CAM IL", "Health Sheet"]

    def test_overrides_are_applied(self, tmp_path):
        out = build_auto_cam_xlsx(tmp_path / "cam_override.xlsx", cibil_score=800, pan="ZZZZZ9999Z")
        wb = openpyxl.load_workbook(out)
        ws = wb["Elegibilty"]
        assert ws["B1"].value == 800
        ws_sys = wb["SystemCam"]
        assert ws_sys["B3"].value == "ZZZZZ9999Z"


class TestChecklistBuilder:
    def test_creates_file_with_checklist_sheet(self, tmp_path):
        out = build_checklist_xlsx(tmp_path / "checklist.xlsx")
        assert out.exists()
        assert out.stat().st_size > 0
        wb = openpyxl.load_workbook(out)
        assert "Checklist" in wb.sheetnames

    def test_no_keys_appear_as_no(self, tmp_path):
        out = build_checklist_xlsx(tmp_path / "checklist_no.xlsx", no_keys=["PAN Card"])
        wb = openpyxl.load_workbook(out)
        ws = wb["Checklist"]
        found_no = any(
            ws.cell(row=r, column=3).value == "No"
            for r in range(1, ws.max_row + 1)
            if ws.cell(row=r, column=2).value == "PAN Card"
        )
        assert found_no

    def test_yes_keys_is_explicit_allow_list(self, tmp_path):
        out = build_checklist_xlsx(tmp_path / "checklist_yes.xlsx", yes_keys=["PAN Card"])
        wb = openpyxl.load_workbook(out)
        ws = wb["Checklist"]
        rows = [
            (ws.cell(row=r, column=2).value, ws.cell(row=r, column=3).value)
            for r in range(2, ws.max_row + 1)
            if ws.cell(row=r, column=2).value
        ]
        statuses = dict(rows)
        assert statuses["PAN Card"] == "Yes"
        assert statuses["Aadhaar Card"] == "NA"


class TestPdSheetBuilder:
    def test_creates_docx_with_table(self, tmp_path):
        out = build_pd_sheet_docx(tmp_path / "pd_sheet.docx")
        assert out.exists()
        assert out.stat().st_size > 0
        doc = Document(out)
        assert len(doc.tables) == 1

    def test_custom_fields_appear_in_table(self, tmp_path):
        fields = {"Business Name": "Seema Kirana", "Annual Revenue": "960000"}
        out = build_pd_sheet_docx(tmp_path / "pd_custom.docx", fields=fields)
        doc = Document(out)
        table = doc.tables[0]
        cell_texts = [cell.text for row in table.rows for cell in row.cells]
        assert "Business Name" in cell_texts
        assert "960000" in cell_texts


class TestEquifaxBuilder:
    def test_creates_html_file(self, tmp_path):
        out = build_equifax_html(tmp_path / "equifax.html")
        assert out.exists()
        assert out.stat().st_size > 0

    def test_credit_score_element_present(self, tmp_path):
        out = build_equifax_html(tmp_path / "equifax_score.html", score=742)
        soup = BeautifulSoup(out.read_text(encoding="utf-8"), "html.parser")
        score_el = soup.find(class_="CreditScore")
        assert score_el is not None
        assert score_el.text.strip() == "742"

    def test_accounts_table_has_rows(self, tmp_path):
        accounts = [
            {
                "lender": "Test Bank",
                "type": "HL",
                "opened": "2022-01",
                "balance": 500000,
                "status": "Active",
            }
        ]
        out = build_equifax_html(tmp_path / "eq_accts.html", accounts=accounts)
        soup = BeautifulSoup(out.read_text(encoding="utf-8"), "html.parser")
        table = soup.find(id="AccountsTable")
        assert table is not None
        rows = table.find_all("tr")
        assert len(rows) >= 2  # header + at least one data row


class TestBankStatementBuilder:
    def test_creates_pdf(self, tmp_path):
        out = build_bank_statement_pdf(tmp_path / "bank.pdf")
        assert out.exists()
        assert out.stat().st_size > 0
        # PDF magic bytes
        assert out.read_bytes()[:4] == b"%PDF"

    def test_custom_account_holder(self, tmp_path):
        import pdfplumber

        out = build_bank_statement_pdf(tmp_path / "bank_custom.pdf", account_holder="RAMESH KUMAR")
        assert out.exists()
        assert out.stat().st_size > 0
        with pdfplumber.open(out) as pdf:
            text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        assert "RAMESH KUMAR" in text


class TestDedupeBuilder:
    def test_creates_file_with_dedupe_sheet(self, tmp_path):
        out = build_dedupe_xlsx(tmp_path / "dedupe.xlsx")
        assert out.exists()
        assert out.stat().st_size > 0
        wb = openpyxl.load_workbook(out)
        assert "Customer_Dedupe" in wb.sheetnames

    def test_header_columns_present(self, tmp_path):
        out = build_dedupe_xlsx(tmp_path / "dedupe_hdr.xlsx")
        wb = openpyxl.load_workbook(out)
        ws = wb["Customer_Dedupe"]
        headers = [ws.cell(row=1, column=c).value for c in range(1, 7)]
        assert headers == ["Customer Name", "Aadhaar", "PAN", "Mobile", "DOB", "Address"]

    def test_custom_customers(self, tmp_path):
        customers = [
            {
                "Customer Name": "TEST PERSON",
                "Aadhaar": "0000-0000-0000",
                "PAN": "TTTTT0000T",
                "Mobile": "0000000000",
                "DOB": "01/01/2000",
                "Address": "Test",
            }
        ]
        out = build_dedupe_xlsx(tmp_path / "dedupe_cust.xlsx", customers=customers)
        wb = openpyxl.load_workbook(out)
        ws = wb["Customer_Dedupe"]
        assert ws.cell(row=2, column=1).value == "TEST PERSON"


class TestCaseZipBuilder:
    def test_creates_zip(self, tmp_path):
        out = build_case_zip(tmp_path / "case.zip")
        assert out.exists()
        assert out.stat().st_size > 0
        assert zipfile.is_zipfile(out)

    def test_zip_has_expected_top_level_folders(self, tmp_path):
        loan_id = "10006484"
        out = build_case_zip(tmp_path / "case_folders.zip", loan_id=loan_id)
        with zipfile.ZipFile(out) as zf:
            names = zf.namelist()
        prefixes = {n.split("/")[0] for n in names}
        assert f"{loan_id}_OTH" in prefixes
        assert f"{loan_id}_BUSINESS_PREMISES" in prefixes
        assert f"{loan_id}_HOUSE_VISIT" in prefixes

    def test_zip_contains_required_files(self, tmp_path):
        loan_id = "10006484"
        out = build_case_zip(tmp_path / "case_files.zip", loan_id=loan_id)
        with zipfile.ZipFile(out) as zf:
            names = set(zf.namelist())
        oth = f"{loan_id}_OTH"
        assert f"{oth}/AUTO_CAM-{loan_id}.xlsx" in names
        assert f"{oth}/Checklist_-{loan_id}.xlsx" in names
        assert f"{oth}/PD_Sheet.docx" in names
        assert f"{oth}/EQUIFAX_CREDIT_REPORT.html" in names
        assert f"{oth}/BANK_STATEMENT_(1).pdf" in names

    def test_zip_has_3_plus_business_premises_photos(self, tmp_path):
        loan_id = "10006484"
        out = build_case_zip(tmp_path / "case_bp.zip", loan_id=loan_id)
        with zipfile.ZipFile(out) as zf:
            bp_photos = [
                n
                for n in zf.namelist()
                if f"{loan_id}_BUSINESS_PREMISES" in n and n.endswith(".jpeg")
            ]
        assert len(bp_photos) >= 3

    def test_zip_has_3_plus_house_visit_photos(self, tmp_path):
        loan_id = "10006484"
        out = build_case_zip(tmp_path / "case_hv.zip", loan_id=loan_id)
        with zipfile.ZipFile(out) as zf:
            hv_photos = [
                n for n in zf.namelist() if f"{loan_id}_HOUSE_VISIT" in n and n.endswith(".jpeg")
            ]
        assert len(hv_photos) >= 3
