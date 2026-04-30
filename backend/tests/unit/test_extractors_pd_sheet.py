"""Tests for PDSheetExtractor — happy path and degraded paths."""

from pathlib import Path

from docx import Document

from app.enums import ExtractionStatus
from app.worker.extractors.pd_sheet import PDSheetExtractor
from tests.fixtures.builders.pd_sheet_builder import build_pd_sheet_docx

# ---------------------------------------------------------------------------
# Happy path
# ---------------------------------------------------------------------------


async def test_pd_sheet_extracts_fields_happy_path(tmp_path: Path):
    p = build_pd_sheet_docx(tmp_path / "pd.docx")
    result = await PDSheetExtractor().extract("pd.docx", p.read_bytes())

    assert result.status == ExtractionStatus.SUCCESS
    assert result.schema_version == "1.0"
    assert result.warnings == []
    assert result.error_message is None

    fields = result.data["fields"]
    assert fields["applicant_name"] == "SEEMA DEVI"
    assert fields["date_of_birth"] == "15/03/1985"
    assert fields["fathers_name"] == "RAMESH KUMAR"
    assert fields["business_type"] == "Kirana Store"
    assert fields["loan_purpose"] == "Business expansion"
    assert fields["existing_loans"] == "None"


async def test_pd_sheet_tables_raw_data_present(tmp_path: Path):
    p = build_pd_sheet_docx(tmp_path / "pd.docx")
    result = await PDSheetExtractor().extract("pd.docx", p.read_bytes())

    assert result.status == ExtractionStatus.SUCCESS
    # tables is a list of 2-D lists
    assert isinstance(result.data["tables"], list)
    assert len(result.data["tables"]) >= 1
    # First row is the header
    assert result.data["tables"][0][0] == ["Question", "Answer"]


async def test_pd_sheet_paragraphs_present(tmp_path: Path):
    p = build_pd_sheet_docx(tmp_path / "pd.docx")
    result = await PDSheetExtractor().extract("pd.docx", p.read_bytes())

    paragraphs = result.data["paragraphs"]
    assert isinstance(paragraphs, list)
    assert any("Personal Discussion" in para for para in paragraphs)


async def test_pd_sheet_extractor_name_and_schema_version():
    extractor = PDSheetExtractor()
    assert extractor.extractor_name == "pd_sheet"
    assert extractor.schema_version == "1.0"


# ---------------------------------------------------------------------------
# Degraded: corrupt bytes → FAILED
# ---------------------------------------------------------------------------


async def test_pd_sheet_failed_on_corrupt_bytes():
    result = await PDSheetExtractor().extract("bad.docx", b"not a docx")

    assert result.status == ExtractionStatus.FAILED
    assert result.error_message is not None
    assert "bad.docx" in result.error_message


# ---------------------------------------------------------------------------
# Degraded: no known fields in table → PARTIAL
# ---------------------------------------------------------------------------


async def test_pd_sheet_partial_when_no_known_fields(tmp_path: Path):
    """Table present but none of the questions match known labels → PARTIAL."""
    p = build_pd_sheet_docx(
        tmp_path / "unknown.docx",
        fields={"Unknown Question 1": "Some Answer", "Unrecognised Field": "Value"},
    )
    result = await PDSheetExtractor().extract("unknown.docx", p.read_bytes())

    assert result.status == ExtractionStatus.PARTIAL
    assert "no_known_fields_matched" in result.warnings


# ---------------------------------------------------------------------------
# Degraded: empty docx (no tables, no paragraphs) → FAILED
# ---------------------------------------------------------------------------


async def test_pd_sheet_failed_on_empty_docx(tmp_path: Path):
    doc = Document()
    out = tmp_path / "empty.docx"
    doc.save(out)

    result = await PDSheetExtractor().extract("empty.docx", out.read_bytes())

    assert result.status == ExtractionStatus.FAILED
    assert result.error_message is not None


# ---------------------------------------------------------------------------
# Edge: label normalisation — mixed case and extra whitespace
# ---------------------------------------------------------------------------


async def test_pd_sheet_normalises_label_case(tmp_path: Path):
    p = build_pd_sheet_docx(
        tmp_path / "case.docx",
        fields={
            "  APPLICANT  NAME  ": "MR TEST",
            "LOAN PURPOSE": "Home Renovation",
        },
    )
    result = await PDSheetExtractor().extract("case.docx", p.read_bytes())

    assert result.status == ExtractionStatus.SUCCESS
    assert result.data["fields"]["applicant_name"] == "MR TEST"
    assert result.data["fields"]["loan_purpose"] == "Home Renovation"


# ---------------------------------------------------------------------------
# Edge: custom fields provided to builder
# ---------------------------------------------------------------------------


async def test_pd_sheet_partial_paragraphs_only_no_tables(tmp_path: Path):
    """Docx with paragraphs but no tables → PARTIAL (no_known_fields_matched)."""
    doc = Document()
    doc.add_paragraph("This document has some text but no Q/A table.")
    doc.add_paragraph("Another paragraph.")
    out = tmp_path / "paragraphs_only.docx"
    doc.save(out)

    result = await PDSheetExtractor().extract("paragraphs_only.docx", out.read_bytes())

    assert result.status == ExtractionStatus.PARTIAL
    assert "no_known_fields_matched" in result.warnings
    assert result.data["tables"] == []
    assert len(result.data["paragraphs"]) >= 1


async def test_pd_sheet_custom_fields(tmp_path: Path):
    p = build_pd_sheet_docx(
        tmp_path / "custom.docx",
        fields={
            "Applicant Name": "JOHN DOE",
            "Date of Birth": "01/01/1990",
            "Address": "456 Park St",
        },
    )
    result = await PDSheetExtractor().extract("custom.docx", p.read_bytes())

    assert result.status == ExtractionStatus.SUCCESS
    assert result.data["fields"]["applicant_name"] == "JOHN DOE"
    assert result.data["fields"]["date_of_birth"] == "01/01/1990"
    assert result.data["fields"]["address"] == "456 Park St"
    # Unset fields remain None
    assert result.data["fields"]["fathers_name"] is None
