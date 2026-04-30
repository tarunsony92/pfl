"""Builder for PD Sheet docx fixture."""

from pathlib import Path

from docx import Document

_DEFAULT_FIELDS = {
    "Applicant Name": "SEEMA DEVI",
    "Date of Birth": "15/03/1985",
    "Father's Name": "RAMESH KUMAR",
    "Address": "123 MG Road, Delhi - 110001",
    "Business Type": "Kirana Store",
    "Years in Business": "7",
    "Monthly Turnover": "80000",
    "Loan Purpose": "Business expansion",
    "Existing Loans": "None",
    "References": "Neighbour - Suresh, 9876543210",
}


def build_pd_sheet_docx(path: Path, fields: dict[str, str] | None = None) -> Path:
    """Create a minimal PD Sheet docx with a Q/A table.

    fields is a dict of question -> answer; defaults are used if not provided.
    Returns path.
    """
    if fields is None:
        fields = _DEFAULT_FIELDS

    doc = Document()
    doc.add_heading("Personal Discussion (PD) Sheet", level=1)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Question"
    hdr[1].text = "Answer"

    for question, answer in fields.items():
        row = table.add_row().cells
        row[0].text = str(question)
        row[1].text = str(answer)

    doc.save(path)
    return path
