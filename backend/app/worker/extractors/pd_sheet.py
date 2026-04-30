"""PDSheetExtractor — parses a Personal Discussion Sheet docx.

Real-world PD docs come in two shapes:
1. Structured table with "Label | Value" rows (builder fixture).
2. Plain paragraph narrative with "Label: Value" lines per paragraph
   (Saksham's actual operations PD — 65 paragraphs, 0 tables).

We scan both tables and paragraphs, use an exact label map plus a fuzzy
substring pass to tolerate wording drift (e.g. "Customer Profile" vs
"Applicant Name", "Business Vintage" vs "Years in Business").
"""

from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from app.enums import ExtractionStatus
from app.worker.extractors.base import BaseExtractor, ExtractionResult

# Exact (normalised) label → field key
_LABEL_MAP: dict[str, str] = {
    "applicant name": "applicant_name",
    "customer profile": "applicant_name",
    "customer name": "applicant_name",
    "date of birth": "date_of_birth",
    "dob": "date_of_birth",
    "father's name": "fathers_name",
    "fathers name": "fathers_name",
    "address": "address",
    "residence address": "address",
    "business type": "business_type",
    "business name": "business_name",
    "years in business": "years_in_business",
    "business vintage": "years_in_business",
    "monthly turnover": "monthly_turnover",
    "monthly income": "monthly_income",
    "monthly income from store": "monthly_income",
    "loan purpose": "loan_purpose",
    "loan amount": "loan_amount",
    "existing loans": "existing_loans",
    "references": "references",
    "operations": "operations",
    "applicant summary": "applicant_summary",
    "field observations": "field_observations",
    "distance from branch": "distance_from_branch",
}

# Fuzzy patterns (substring on normalised label). Order: most specific first.
_FUZZY_PATTERNS: list[tuple[str, str]] = [
    ("monthly income", "monthly_income"),
    ("monthly turnover", "monthly_turnover"),
    ("business vintage", "years_in_business"),
    ("years in business", "years_in_business"),
    ("business type", "business_type"),
    ("business name", "business_name"),
    ("loan purpose", "loan_purpose"),
    ("loan amount", "loan_amount"),
    ("customer profile", "applicant_name"),
    ("applicant name", "applicant_name"),
    ("customer name", "applicant_name"),
    ("father", "fathers_name"),
]


def _normalise(text: str) -> str:
    return " ".join(text.strip().split()).lower()


def _resolve_label(raw_label: str) -> str | None:
    """Return the canonical field key for a label, or None if nothing matches."""
    norm = _normalise(raw_label).rstrip(":").strip()
    if norm in _LABEL_MAP:
        return _LABEL_MAP[norm]
    for pattern, key in _FUZZY_PATTERNS:
        if pattern in norm:
            return key
    return None


class PDSheetExtractor(BaseExtractor):
    extractor_name = "pd_sheet"
    schema_version = "1.0"

    async def extract(self, filename: str, body_bytes: bytes) -> ExtractionResult:
        try:
            doc = Document(BytesIO(body_bytes))
        except (PackageNotFoundError, Exception) as exc:
            return ExtractionResult(
                status=ExtractionStatus.FAILED,
                schema_version=type(self).schema_version,
                data={},
                error_message=f"python-docx failed to open {filename!r}: {exc}",
            )

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Extract all tables as 2-D lists
        raw_tables: list[list[list[str]]] = []
        for tbl in doc.tables:
            raw_tables.append([[cell.text for cell in row.cells] for row in tbl.rows])

        if not raw_tables and not paragraphs:
            return ExtractionResult(
                status=ExtractionStatus.FAILED,
                schema_version=type(self).schema_version,
                data={},
                error_message="Document contains no tables and no paragraphs.",
            )

        # Superset of all values from _LABEL_MAP + fuzzy map so the dict keys
        # always include every recognised canonical field.
        all_field_keys = set(_LABEL_MAP.values()) | {v for _, v in _FUZZY_PATTERNS}
        fields: dict[str, Any] = {k: None for k in all_field_keys}

        # Pass 1: any table with 2+ columns — iterate rows, match col[0]→label,
        # col[1]→value. Skip a likely header row (both cells look like labels).
        for tbl in raw_tables:
            for row in tbl:
                if len(row) < 2:
                    continue
                field_key = _resolve_label(row[0])
                if field_key is not None and fields.get(field_key) is None:
                    value = row[1].strip()
                    if value:
                        fields[field_key] = value

        # Pass 2: paragraphs like "Label: Value" (real narrative PD).
        for para in paragraphs:
            if ":" not in para:
                continue
            label_part, _, value_part = para.partition(":")
            value = value_part.strip()
            if not value:
                continue
            field_key = _resolve_label(label_part)
            if field_key is not None and fields.get(field_key) is None:
                fields[field_key] = value

        data: dict[str, Any] = {
            "fields": fields,
            "tables": raw_tables,
            "paragraphs": paragraphs,
        }

        # Determine status — primary output is ≥1 known field matched.
        known_matched = any(v is not None for v in fields.values())
        warnings: list[str] = []
        if known_matched:
            status = ExtractionStatus.SUCCESS
        else:
            warnings.append("no_known_fields_matched")
            status = ExtractionStatus.PARTIAL

        return ExtractionResult(
            status=status,
            schema_version=type(self).schema_version,
            data=data,
            warnings=warnings,
        )
